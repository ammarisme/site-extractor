from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.common.exceptions import TimeoutException, WebDriverException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from urllib.parse import urljoin
from docx import Document
import os, sys, uuid
from contextlib import asynccontextmanager

# Directory to save extracted documents
OUTPUT_DIR = "extracted_documents"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Lifespan handler replaces on_event startup/shutdown
@asynccontextmanager
async def lifespan(app: FastAPI):
    # --- startup logic ---
    options = webdriver.ChromeOptions()
    options.add_argument("--headless")
    options.add_argument("--disable-gpu")
    options.add_argument("--log-level=3")  # suppress verbose logs
    options.add_experimental_option("excludeSwitches", ["enable-logging"])
    service = Service(ChromeDriverManager().install())
    app.state.driver = webdriver.Chrome(service=service, options=options)
    yield
    # --- shutdown logic ---
    app.state.driver.quit()

app = FastAPI(lifespan=lifespan)

class ExtractRequest(BaseModel):
    start_url: str
    prefix: str


def extract_links_with_prefix(driver, url: str, prefix: str, timeout: int = 10):
    """
    Navigate to `url` and return all <a> hrefs that start with `prefix`.
    """
    try:
        driver.get(url)
        WebDriverWait(driver, timeout).until(
            EC.presence_of_all_elements_located((By.TAG_NAME, "a"))
        )
    except (TimeoutException, WebDriverException) as e:
        print(f"Failed to load {url}: {e}", file=sys.stderr)
        return []

    links = []
    for elem in driver.find_elements(By.TAG_NAME, "a"):
        href = elem.get_attribute("href")
        if not href:
            continue
        absolute_href = urljoin(url, href)
        if absolute_href.startswith(prefix):
            links.append(absolute_href)
    return links


def extract_page_text(driver, url: str, timeout: int = 10):
    """
    Navigate to `url` and return text of the main content area.
    """
    try:
        driver.get(url)
        WebDriverWait(driver, timeout).until(
            EC.presence_of_element_located((By.ID, "content-area"))
        )
        content_elem = driver.find_element(By.ID, "content-area")
        return content_elem.text.strip()
    except (TimeoutException, WebDriverException) as e:
        print(f"Failed to fetch or parse {url}: {e}", file=sys.stderr)
        return ""

@app.post("/extract")
def extract_site(request: ExtractRequest):
    """
    Extracts all pages linked from `start_url` with links prefixed by `prefix`,
    compiles them into a single .docx, and saves under `extracted_documents/`.
    Returns the saved file path.
    """
    driver = app.state.driver
    print(f"Gathering links from {request.start_url}…")
    links = extract_links_with_prefix(driver, request.start_url, request.prefix)
    if not links:
        print("No matching links found.")
        raise HTTPException(status_code=404, detail="No matching links found.")

    doc = Document()
    for link in links:
        print(f"Extracting content from: {link}")
        text = extract_page_text(driver, link)
        if text:
            doc.add_heading(link, level=1)
            doc.add_paragraph(text)
        else:
            print(f"No content extracted from {link}", file=sys.stderr)

    file_id = uuid.uuid4().hex
    filename = f"{file_id}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    try:
        doc.save(output_path)
        print(f"Document saved as {output_path}")
    except Exception as e:
        print(f"Failed to save document: {e}", file=sys.stderr)
        raise HTTPException(status_code=500, detail=f"Could not save document: {e}")

    return {"file_path": output_path}

# Allow `python server.py` to launch Uvicorn directly
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("server:app", host="127.0.0.1", port=8000, reload=True)
