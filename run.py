from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.common.exceptions import TimeoutException, WebDriverException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from urllib.parse import urljoin
from docx import Document
import sys

# Configure Selenium to use Chrome in headless mode,
# with automatic driver management via webdriver-manager
options = webdriver.ChromeOptions()
options.add_argument("--headless")
options.add_argument("--disable-gpu")
service = Service(ChromeDriverManager().install())
driver = webdriver.Chrome(service=service, options=options)

def extract_links_with_prefix(driver, url, prefix, timeout=10):
    """Navigate to `url` and return all <a> hrefs that start with `prefix`."""
    try:
        driver.get(url)
        WebDriverWait(driver, timeout).until(
            EC.presence_of_all_elements_located((By.TAG_NAME, "a"))
        )
    except (TimeoutException, WebDriverException) as e:
        print(f"Failed to load {url}: {e}", file=sys.stderr)
        return []

    links = []
    for elem in driver.find_elements(By.TAG_NAME, "a"):
        href = elem.get_attribute("href")
        if not href:
            continue
        absolute_href = urljoin(url, href)
        if absolute_href.startswith(prefix):
            links.append(absolute_href)
    return links

def extract_page_text(driver, url, timeout=10):
    """Navigate to `url` and return text of the main content area."""
    try:
        driver.get(url)
        WebDriverWait(driver, timeout).until(
            EC.presence_of_element_located((By.ID, "content-area"))
        )
        content = driver.find_element(By.ID, "content-area").text.strip()
        return content
    except (TimeoutException, WebDriverException) as e:
        print(f"Failed to fetch or parse {url}: {e}", file=sys.stderr)
        return ""

def main():
    start_url = "https://docs.pipecat.ai/guides/introduction"
    prefix = "https://docs.pipecat.ai/"

    print(f"Gathering links from {start_url}…")
    links = extract_links_with_prefix(driver, start_url, prefix)
    if not links:
        print("No matching links found.")
        return

    doc = Document()

    for link in links:
        print(f"Extracting content from: {link}")
        content = extract_page_text(driver, link)
        if content:
            doc.add_heading(link, level=1)
            doc.add_paragraph(content)
        else:
            print(f"No content extracted from {link}", file=sys.stderr)

    output_path = "merged_content.docx"
    doc.save(output_path)
    print(f"Document saved as {output_path}")

main()
# if __name__ == "__main__":
#     try:
#         main()
#     finally:
#         driver.quit()
